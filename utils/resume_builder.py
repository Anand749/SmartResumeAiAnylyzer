from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import tempfile
import traceback
import re

class ResumeBuilder:
    """A class to generate professional resumes in various templates with validation."""
    
    def __init__(self):
        """Initialize the ResumeBuilder with available templates and requirements."""
        self.templates = {
            "modern": self.build_modern_template,
            "professional": self.build_professional_template,
            "minimal": self.build_minimal_template,
            "creative": self.build_creative_template
        }
        
        self.required_fields = {
            'personal_info': ['full_name', 'email', 'phone'],
            'sections': ['experience', 'education']
        }

    def validate_email(self, email):
        """Validate email format using regex.
        
        Args:
            email (str): The email address to validate
            
        Returns:
            bool: True if email is valid, False otherwise
        """
        if not email:
            return False
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return re.match(pattern, email) is not None
    
    def validate_phone(self, phone):
        """Basic phone number validation.
        
        Args:
            phone (str): The phone number to validate
            
        Returns:
            bool: True if phone has at least 7 digits, False otherwise
        """
        if not phone:
            return False
        digits = re.sub(r'[^0-9]', '', phone)
        return len(digits) >= 7
    
    def validate_data(self, data):
        """Validate resume data before generation.
        
        Args:
            data (dict): The resume data to validate
            
        Returns:
            list: List of error messages, empty if valid
        """
        errors = []
        
        if not isinstance(data, dict):
            return ["Invalid data format. Expected dictionary."]
        
        # Check template is valid
        template = data.get('template', 'modern').lower()
        if template not in self.templates:
            errors.append(f"Invalid template '{template}'. Available options: {', '.join(self.templates.keys())}")
        
        # Check personal info exists
        if 'personal_info' not in data:
            errors.append("Missing personal information section")
            return errors
            
        personal_info = data['personal_info']
        
        # Validate required fields
        for field in self.required_fields['personal_info']:
            if field not in personal_info or not str(personal_info[field]).strip():
                errors.append(f"Missing required field: {field}")
            elif field == 'email' and not self.validate_email(personal_info[field]):
                errors.append("Invalid email format (example: user@example.com)")
            elif field == 'phone' and not self.validate_phone(personal_info[field]):
                errors.append("Phone number should have at least 7 digits")
        
        # Validate at least one required section exists with valid data
        has_required_section = False
        for section in self.required_fields['sections']:
            if section in data and data[section]:
                if isinstance(data[section], list) and len(data[section]) > 0:
                    has_required_section = True
                    # Validate each item in the section
                    for i, item in enumerate(data[section]):
                        if not isinstance(item, dict):
                            errors.append(f"Invalid format in {section} item {i+1}")
                            continue
                            
                        key = 'company' if section == 'experience' else 'school'
                        if not item.get(key, '').strip():
                            errors.append(f"Missing {key} in {section} item {i+1}")
                            
                        # Validate dates if present
                        for date_field in ['start_date', 'end_date']:
                            if date_field in item and not item[date_field]:
                                errors.append(f"Missing {date_field} in {section} item {i+1}")
        
        if not has_required_section:
            errors.append("Please add at least one experience or education entry")
        
        return errors

    def generate_resume(self, data):
        """Generate a resume based on the provided data and template.
        
        Args:
            data (dict): Resume data including template choice
            
        Returns:
            BytesIO: In-memory file containing the generated resume
            
        Raises:
            ValueError: If validation fails or data is incomplete
            Exception: For other generation errors
        """
        try:
            # Validate data before generation
            validation_errors = self.validate_data(data)
            if validation_errors:
                raise ValueError("\n".join([f"• {error}" for error in validation_errors]))
            
            print(f"Starting resume generation with template: {data.get('template', 'modern')}")
            
            # Create a new document
            doc = Document()
            
            # Select and apply template (default to modern if not specified)
            template_name = data.get('template', 'modern').lower()
            template_func = self.templates.get(template_name, self.build_modern_template)
            print(f"Using template: {template_name}")
            
            doc = template_func(doc, data)
            
            # Save to buffer
            buffer = BytesIO()
            print("Saving document to buffer...")
            doc.save(buffer)
            buffer.seek(0)
            print("Resume generated successfully!")
                
            return buffer
            
        except ValueError as ve:
            print(f"Validation failed: {str(ve)}")
            raise ValueError(f"Please fix these issues:\n{str(ve)}") from None
        except Exception as e:
            print(f"Error in generate_resume: {str(e)}")
            print(f"Full traceback: {traceback.format_exc()}")
            raise RuntimeError("Failed to generate resume. Please try again or contact support.") from e

    # ===== TEMPLATE METHODS =====
    def build_modern_template(self, doc, data):
        """Build modern style resume with clean, minimalist design."""
        try:
            # Add header with name and contact info
            header = doc.sections[0].header
            header_table = header.add_table(1, 3, Inches(6))
            header_table.autofit = False
            
            # Set column widths
            hdr_cells = header_table.rows[0].cells
            hdr_cells[0].width = Inches(2)
            hdr_cells[1].width = Inches(2)
            hdr_cells[2].width = Inches(2)
            
            # Add name
            name = hdr_cells[0].paragraphs[0]
            name_run = name.add_run(data['personal_info']['full_name'])
            name_run.font.size = Pt(16)
            name_run.bold = True
            
            # Add contact info
            contact = hdr_cells[2].paragraphs[0]
            contact_run = contact.add_run(
                f"{data['personal_info']['email']}\n"
                f"{data['personal_info']['phone']}"
            )
            contact_run.font.size = Pt(10)
            contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add sections
            self._add_section(doc, "PROFESSIONAL SUMMARY", data.get('summary', ''))
            self._add_experience_section(doc, data.get('experience', []))
            self._add_education_section(doc, data.get('education', []))
            self._add_skills_section(doc, data.get('skills', []))
            
            return doc
        except Exception as e:
            print(f"Error in build_modern_template: {str(e)}")
            raise

    def build_professional_template(self, doc, data):
        """Build professional style resume with formal layout."""
        try:
            # Add title with name
            title = doc.add_paragraph()
            title_run = title.add_run(data['personal_info']['full_name'])
            title_run.font.size = Pt(22)
            title_run.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add contact info below name
            contact = doc.add_paragraph()
            contact_run = contact.add_run(
                f"{data['personal_info']['email']} | "
                f"{data['personal_info']['phone']}"
            )
            contact_run.font.size = Pt(11)
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add horizontal line
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            self._add_horizontal_line(doc)
            
            # Add sections
            self._add_section(doc, "SUMMARY", data.get('summary', ''))
            self._add_experience_section(doc, data.get('experience', []))
            self._add_education_section(doc, data.get('education', []))
            self._add_skills_section(doc, data.get('skills', []))
            
            return doc
        except Exception as e:
            print(f"Error in build_professional_template: {str(e)}")
            raise

    def build_minimal_template(self, doc, data):
        """Build minimal style resume with essential information only."""
        try:
            # Add name
            name = doc.add_paragraph()
            name_run = name.add_run(data['personal_info']['full_name'])
            name_run.font.size = Pt(14)
            name_run.bold = True
            
            # Add contact info
            doc.add_paragraph(data['personal_info']['email'])
            doc.add_paragraph(data['personal_info']['phone'])
            doc.add_paragraph()
            
            # Add sections with minimal formatting
            if data.get('experience'):
                doc.add_paragraph("EXPERIENCE", style='Heading2')
                for exp in data['experience']:
                    doc.add_paragraph(f"{exp.get('company', '')} | {exp.get('position', '')}")
                    doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                    for desc in self._format_list_items(exp.get('description', [])):
                        doc.add_paragraph(f"• {desc}", style='ListBullet')
                    doc.add_paragraph()
            
            if data.get('education'):
                doc.add_paragraph("EDUCATION", style='Heading2')
                for edu in data['education']:
                    doc.add_paragraph(f"{edu.get('school', '')} | {edu.get('degree', '')}")
                    doc.add_paragraph(edu.get('graduation_date', ''))
                    doc.add_paragraph()
            
            return doc
        except Exception as e:
            print(f"Error in build_minimal_template: {str(e)}")
            raise

    def build_creative_template(self, doc, data):
        """Build creative style resume with unique design elements."""
        try:
            # Add colorful header
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_run = header_para.add_run(data['personal_info']['full_name'])
            header_run.font.size = Pt(24)
            header_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)  # Blue
            header_run.bold = True
            
            # Add contact info with icons (represented by text here)
            contact = doc.add_paragraph()
            contact_run = contact.add_run(
                f"✉ {data['personal_info']['email']}   "
                f"📞 {data['personal_info']['phone']}"
            )
            contact_run.font.size = Pt(11)
            contact_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            
            # Add decorative horizontal line
            self._add_decorative_line(doc)
            
            # Add sections with creative formatting
            self._add_creative_section(doc, "PROFILE", data.get('summary', ''))
            self._add_creative_experience(doc, data.get('experience', []))
            self._add_creative_education(doc, data.get('education', []))
            self._add_creative_skills(doc, data.get('skills', []))
            
            return doc
        except Exception as e:
            print(f"Error in build_creative_template: {str(e)}")
            raise

    # ===== HELPER METHODS =====
    def _add_section(self, doc, title, content):
        """Add a standard section to the document."""
        if not content:
            return
            
        doc.add_paragraph(title, style='Heading2')
        for item in self._format_list_items(content):
            doc.add_paragraph(item)
        doc.add_paragraph()

    def _add_experience_section(self, doc, experiences):
        """Add experience section to the document."""
        if not experiences:
            return
            
        doc.add_paragraph("EXPERIENCE", style='Heading2')
        for exp in experiences:
            # Add company and position
            p = doc.add_paragraph()
            p.add_run(exp.get('company', '')).bold = True
            p.add_run(" | " + exp.get('position', ''))
            
            # Add dates
            doc.add_paragraph(
                f"{exp.get('start_date', '')} - {exp.get('end_date', '')}",
                style='BodyText'
            )
            
            # Add description items
            for desc in self._format_list_items(exp.get('description', [])):
                doc.add_paragraph(f"• {desc}", style='ListBullet')
            
            doc.add_paragraph()

    def _add_education_section(self, doc, educations):
        """Add education section to the document."""
        if not educations:
            return
            
        doc.add_paragraph("EDUCATION", style='Heading2')
        for edu in educations:
            # Add school and degree
            p = doc.add_paragraph()
            p.add_run(edu.get('school', '')).bold = True
            p.add_run(" | " + edu.get('degree', ''))
            
            # Add dates
            doc.add_paragraph(
                edu.get('graduation_date', ''),
                style='BodyText'
            )
            doc.add_paragraph()

    def _add_skills_section(self, doc, skills):
        """Add skills section to the document."""
        if not skills:
            return
            
        doc.add_paragraph("SKILLS", style='Heading2')
        skills_text = ", ".join(self._format_list_items(skills))
        doc.add_paragraph(skills_text)
        doc.add_paragraph()

    def _add_horizontal_line(self, doc):
        """Add a simple horizontal line to the document."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_break()
        p.paragraph_format.border_bottom = True

    def _add_decorative_line(self, doc):
        """Add a decorative horizontal line to the document."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run("━" * 30)  # Using box drawing character
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run.font.size = Pt(12)

    def _add_creative_section(self, doc, title, content):
        """Add a creatively styled section to the document."""
        if not content:
            return
            
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        run.bold = True
        run.font.size = Pt(14)
        
        for item in self._format_list_items(content):
            doc.add_paragraph(item)
        doc.add_paragraph()

    def _add_creative_experience(self, doc, experiences):
        """Add creatively styled experience section."""
        if not experiences:
            return
            
        self._add_creative_section(doc, "EXPERIENCE", "")
        for exp in experiences:
            # Add company and position
            p = doc.add_paragraph()
            run = p.add_run(exp.get('company', ''))
            run.bold = True
            run.font.size = Pt(12)
            p.add_run(" ― " + exp.get('position', ''))
            
            # Add dates
            doc.add_paragraph(
                f"{exp.get('start_date', '')} → {exp.get('end_date', '')}",
                style='BodyText'
            )
            
            # Add description items with custom bullet points
            for desc in self._format_list_items(exp.get('description', [])):
                p = doc.add_paragraph()
                p.add_run("◦ ").font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
                p.add_run(desc)
            
            doc.add_paragraph()

    def _add_creative_education(self, doc, educations):
        """Add creatively styled education section."""
        if not educations:
            return
            
        self._add_creative_section(doc, "EDUCATION", "")
        for edu in educations:
            # Add school and degree
            p = doc.add_paragraph()
            run = p.add_run(edu.get('school', ''))
            run.bold = True
            run.font.size = Pt(12)
            p.add_run(" ― " + edu.get('degree', ''))
            
            # Add dates
            doc.add_paragraph(
                edu.get('graduation_date', ''),
                style='BodyText'
            )
            doc.add_paragraph()

    def _add_creative_skills(self, doc, skills):
        """Add creatively styled skills section."""
        if not skills:
            return
            
        self._add_creative_section(doc, "SKILLS", "")
        skills_text = "  •  ".join(self._format_list_items(skills))
        p = doc.add_paragraph()
        run = p.add_run(skills_text)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        doc.add_paragraph()

    def _format_list_items(self, items):
        """Helper function to handle both string and list inputs.
        
        Args:
            items (str or list): Input items to format
            
        Returns:
            list: Cleaned list of items
        """
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []